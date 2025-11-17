#!/usr/bin/env python3
"""
Create nav.docx and footer.docx files with proper table structures for Google Docs upload.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_nav_doc():
    """Create nav.docx with 3-row table structure."""
    doc = Document()
    
    # Create 3x1 table
    table = doc.add_table(rows=3, cols=1)
    table.style = 'Table Grid'
    
    # Row 1: Brand
    row1_cell = table.rows[0].cells[0]
    row1_para = row1_cell.paragraphs[0]
    row1_para.text = 'Acerbis'
    row1_para.style = 'Normal'
    
    # Row 2: Navigation Sections with bullets
    row2_cell = table.rows[1].cells[0]
    row2_cell.text = ''  # Clear default text
    
    nav_items = [
        ('Chi siamo', 1),
        ('R&D e Qualità', 1),
        ('Sostenibilità', 1),
        ('Settori', 0),
        ('Motorcycle', 1),
        ('Agricoltura', 1),
        ('Movimento terra', 1),
        ('Material handling', 1),
        ('Altri settori', 1),
        ('Tecnologie', 0),
        ('Stampaggio rotazionale', 1),
        ('Stampaggio a iniezione', 1),
        ('Engineering', 0),
        ('Project Management', 1),
        ('Engineering Capability', 1),
        ('Lavora con noi', 0),
        ('Contatti', 0),
        ('Mondo Acerbis', 0),
    ]
    
    for i, (text, indent_level) in enumerate(nav_items):
        if i == 0:
            para = row2_cell.paragraphs[0]
        else:
            para = row2_cell.add_paragraph()
        
        para.text = text
        para.style = 'List Bullet' if indent_level == 0 else 'List Bullet 2'
        para.level = indent_level
    
    # Row 3: Tools (empty)
    row3_cell = table.rows[2].cells[0]
    row3_cell.text = ''
    
    # Save
    doc.save('nav.docx')
    print('✅ Created: nav.docx')

def create_footer_doc():
    """Create footer.docx with 1x4 table structure."""
    doc = Document()
    
    # Create 1x4 table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    row = table.rows[0]
    
    # Column 1: Acerbis
    col1_items = [
        'Acerbis',
        'Chi siamo',
        'R&D e Qualità',
        'Sostenibilità',
    ]
    col1_cell = row.cells[0]
    col1_cell.text = ''
    for i, text in enumerate(col1_items):
        if i == 0:
            para = col1_cell.paragraphs[0]
        else:
            para = col1_cell.add_paragraph()
        para.text = text
        para.style = 'List Bullet'
        para.level = 0
    
    # Column 2: Settori
    col2_items = [
        'Settori',
        'Motorcycle',
        'Agricoltura',
        'Movimento terra',
        'Material handling',
        'Altri settori',
    ]
    col2_cell = row.cells[1]
    col2_cell.text = ''
    for i, text in enumerate(col2_items):
        if i == 0:
            para = col2_cell.paragraphs[0]
        else:
            para = col2_cell.add_paragraph()
        para.text = text
        para.style = 'List Bullet'
        para.level = 0
    
    # Column 3: Tecnologie
    col3_items = [
        'Tecnologie',
        'Stampaggio rotazionale',
        'Stampaggio a iniezione',
    ]
    col3_cell = row.cells[2]
    col3_cell.text = ''
    for i, text in enumerate(col3_items):
        if i == 0:
            para = col3_cell.paragraphs[0]
        else:
            para = col3_cell.add_paragraph()
        para.text = text
        para.style = 'List Bullet'
        para.level = 0
    
    # Column 4: Engineering + Copyright
    col4_items = [
        'Engineering',
        'Project Management',
        'Engineering Capability',
        '© 2024 ACERBIS ITALIA SPA',
        'Società soggetta all\'attività di direzione e coordinamento di Yellow Holding S.r.l.',
        'P.IVA 00862020161',
        'Privacy & Cookie Policy',
        'Contatti',
        'Lavora con noi',
    ]
    col4_cell = row.cells[3]
    col4_cell.text = ''
    for i, text in enumerate(col4_items):
        if i == 0:
            para = col4_cell.paragraphs[0]
        else:
            para = col4_cell.add_paragraph()
        para.text = text
        para.style = 'List Bullet'
        para.level = 0
    
    # Save
    doc.save('footer.docx')
    print('✅ Created: footer.docx')

if __name__ == '__main__':
    print('Creating nav.docx and footer.docx...')
    create_nav_doc()
    create_footer_doc()
    print('\n✅ Done! Files created:')
    print('   - nav.docx')
    print('   - footer.docx')
    print('\nNext steps:')
    print('1. Download these files')
    print('2. Upload to Google Drive folder: https://drive.google.com/drive/folders/1KcRM48-YYJall7YnzW_Qg66PPLWQLAns')
    print('3. Rename to "nav" and "footer" (remove .docx)')
    print('4. Share with helix@adobe.com (Editor)')
    print('5. Publish via Sidekick')
